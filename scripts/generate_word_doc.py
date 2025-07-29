import os

from docx.shared import Inches
from docx import Document

def create_doc(path, df_stats, df_stats_by_theta, motion, atlas, repl, mc, theta_results, fig_dir=None):
    """
    df_stats: a dict or DataFrame containing:
      - n_subjects
      - mean_S_awake, sd_S_awake
      - mean_S_deep,  sd_S_deep
      - t_S,   p_S,   d_S
      - mean_CI_awake, sd_CI_awake
      - mean_CI_deep,  sd_CI_deep
      - t_CI,  p_CI,  d_CI
    motion: dict of your motion regression results
    atlas: dict of robustness results
    repl: replication dict
    mc: model comparison dict
    """
    doc = Document()
    doc.add_heading("Empirical Validation of Synergy Metrics in fMRI Data", level=1)

    # 1) Main effects with descriptives
    doc.add_paragraph(
        f"In {df_stats['n_subjects']} subjects, Synergy (S) was higher in the awake state "
        f"({df_stats['mean_S_awake']:.3f} ± {df_stats['sd_S_awake']:.3f}) than in deep sedation "
        f"({df_stats['mean_S_deep']:.3f} ± {df_stats['sd_S_deep']:.3f}; "
        f"t({df_stats['df']})={df_stats['t_S']:.2f}, p={df_stats['p_S']:.4f}, "
        f"Cohen's d={df_stats['d_S']:.2f})."
    )
    doc.add_paragraph(
        f"Consciousness Index (CI) likewise decreased under sedation "
        f"({df_stats['mean_CI_awake']:.3f} ± {df_stats['sd_CI_awake']:.3f} vs. "
        f"{df_stats['mean_CI_deep']:.3f} ± {df_stats['sd_CI_deep']:.3f}; "
        f"t({df_stats['df']})={df_stats['t_CI']:.2f}, p={df_stats['p_CI']:.4f}, "
        f"d={df_stats['d_CI']:.2f})."
    )

    # 2) Motion regressors
    coefs = ", ".join(f"{c:.3f}" for c in motion['coef_awake'])
    pvals = ", ".join(f"{p:.3f}" for p in motion['p_awake'])
    doc.add_paragraph(
        f"Motion covariates (awake): coefs=[{coefs}]; p‐values=[{pvals}]."
    )

    # 3) Atlas robustness
    for name, r in atlas.items():
        doc.add_paragraph(
            f"Atlas {name}: metric roughness awake={r['awake']:.3f}, deep={r['deep']:.3f}."
        )

    # 4) Replication
 #   doc.add_paragraph(
 #       f"Replication ΔS={repl['delta_S']:.3f} "
 #       f"(95% CI [{repl['ci'][0]:.3f}–{repl['ci'][1]:.3f}]), Cohen's d={repl['cohend']:.2f}."
 #   )
     # 5) Per-θ Synergy statistics
    doc.add_heading("Synergy by θ", level=2)
    for theta, row in df_stats_by_theta.iterrows():
        doc.add_paragraph(
            f"θ={theta:.2f}: t_S={row.t_S:.2f}, p_S={row.p_S:.4f}, "
            f"d_S={row.d_S:.2f}"
        )
    
    # 5b) Per-θ permutation test AUCs
    doc.add_heading("Permutation-Test AUC by θ", level=2)
    for theta, res in theta_results.items():
        doc.add_paragraph(
            f"θ={theta:.2f}: AUC_S={res['auc_S']:.3f}, p_S={res['p_S']:.4f}; "
            f"AUC_CI={res['auc_CI']:.3f}, p_CI={res['p_CI']:.4f}"
        )

    # 6) Supplemental θ‐curve figure
    if fig_dir is not None:
        theta_fig = os.path.join(fig_dir, 'supp_theta_curve.png')
        if os.path.exists(theta_fig):
            doc.add_heading("Supplementary: Synergy difference across θ", level=2)
            # Insert the figure at ~5″ wide (adjust as needed)
            doc.add_picture(theta_fig, width=Inches(5.0))
        else:
            doc.add_paragraph(f"(Could not find θ‐curve at {theta_fig})")


    # 7) Model comparisons
    for m, res in mc.items():
        doc.add_paragraph(
            f"Model {m}: ΔAUC={res['delta_auc']:.3f}, p={res['p_val']:.3f}."
        )

    doc.save(path)

