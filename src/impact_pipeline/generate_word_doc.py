import os
import numpy as np
import pandas as pd
from typing import Optional, Union
from docx.shared import Inches
from docx import Document
import matplotlib.pyplot as plt
# ---------- formatting ----------

def fmt(x, decimals=3, sci_below=1e-4, sci_above=1e5):
    """Human-friendly float formatting. Uses scientific notation if |x| is
    very small/large to avoid '0.000' artifacts."""
    if x is None or (isinstance(x, float) and not np.isfinite(x)):
        return "na"
    ax = abs(x)
    if (ax != 0 and ax < sci_below) or ax >= sci_above:
        return f"{x:.3e}"
    return f"{x:.{decimals}f}"

def fmt_pm(mean, err, decimals=3):
    return f"{fmt(mean, decimals)} ± {fmt(err, decimals)}"

def safe_pct(num, den, decimals=1):
    if den is None or not np.isfinite(den) or den == 0:
        return "na"
    return f"{(num/den*100):.{decimals}f}%"

# ---------- stats ----------

def _paired_stats(pivot_a, pivot_b):
    df_pair = pd.DataFrame({'a': pivot_a, 'b': pivot_b}).dropna()
    x = df_pair['a'].values; y = df_pair['b'].values
    n = len(df_pair)
    if n < 2:
        return np.nan, np.nan, np.nan, 0
    diff = x - y
    mean_diff = diff.mean()
    sd_diff = diff.std(ddof=1)
    t = mean_diff / (sd_diff / np.sqrt(n)) if sd_diff > 0 else np.nan
    try:
        from scipy import stats
        p = 2 * stats.t.sf(np.abs(t), df=n-1) if np.isfinite(t) else np.nan
    except Exception:
        p = np.nan
    d = mean_diff / sd_diff if sd_diff > 0 else np.nan
    return float(t), float(p), float(d), n - 1

def _group_descriptives(series):
    n = series.notna().sum()
    mean = series.mean()
    sd = series.std(ddof=1) if n > 1 else np.nan
    sem = sd / np.sqrt(n) if n > 0 and np.isfinite(sd) else np.nan
    return n, mean, sd, sem

def _safe_row_by_theta(df_stats_by_theta, target_theta, tol=1e-6):
    if df_stats_by_theta is None or df_stats_by_theta.empty:
        return None
    idx = df_stats_by_theta.index
    nearest = float(idx[np.argmin(np.abs(idx - target_theta))])
    if abs(nearest - target_theta) <= max(tol, 1e-9):
        return df_stats_by_theta.loc[nearest]
    return df_stats_by_theta.loc[nearest]


def _best_theta_from_stats(df_stats_by_theta: Optional[pd.DataFrame]) -> float:
    if df_stats_by_theta is None or df_stats_by_theta.empty:
        return np.nan
    if "mean_diff_S" not in df_stats_by_theta.columns:
        return np.nan
    vals = df_stats_by_theta["mean_diff_S"].to_numpy(dtype=float)
    if not np.isfinite(vals).any():
        return np.nan
    idx = df_stats_by_theta.index.to_numpy(dtype=float)
    return float(idx[int(np.nanargmax(np.abs(vals)))])


def _save_paired_metric_plot(
    agg: pd.DataFrame,
    metric: str,
    out_path: str,
    title: str,
    y_label: str,
    scale: float = 1.0,
    sessions=("awake", "deep"),
):
    if metric not in agg.columns:
        return False

    piv = agg.pivot(index='subject', columns='session', values=metric)
    if not set(sessions).issubset(set(piv.columns)):
        return False

    paired = piv[[sessions[0], sessions[1]]].dropna()
    if paired.empty:
        return False

    a = paired[sessions[0]].to_numpy(dtype=float) * scale
    d = paired[sessions[1]].to_numpy(dtype=float) * scale
    means = [np.nanmean(a), np.nanmean(d)]
    sems = [
        np.nanstd(a, ddof=1) / np.sqrt(len(a)) if len(a) > 1 else np.nan,
        np.nanstd(d, ddof=1) / np.sqrt(len(d)) if len(d) > 1 else np.nan,
    ]

    x = np.array([0.0, 1.0], dtype=float)
    rng = np.random.RandomState(42)
    jitter = 0.045

    fig, ax = plt.subplots(figsize=(6.0, 4.0))
    for i in range(len(paired)):
        ax.plot(x, [a[i], d[i]], color="#B8B8B8", linewidth=0.7, alpha=0.6, zorder=1)

    ax.errorbar(
        x,
        means,
        yerr=sems,
        fmt='o',
        color="#1f4e79",
        ecolor="#1f4e79",
        elinewidth=1.5,
        capsize=4,
        markersize=7,
        zorder=3,
        label="Mean ± SEM",
    )

    ax.scatter(
        np.full_like(a, x[0]) + rng.uniform(-jitter, jitter, size=len(a)),
        a,
        color="#2a9d8f",
        alpha=0.75,
        s=24,
        zorder=2,
        label=sessions[0] if len(paired) > 0 else None,
    )
    ax.scatter(
        np.full_like(d, x[1]) + rng.uniform(-jitter, jitter, size=len(d)),
        d,
        color="#e76f51",
        alpha=0.75,
        s=24,
        zorder=2,
        label=sessions[1] if len(paired) > 0 else None,
    )

    ax.set_xticks(x)
    ax.set_xticklabels([sessions[0].title(), sessions[1].title()])
    ax.set_ylabel(y_label)
    ax.set_title(title)
    ax.grid(axis='y', linestyle=':', alpha=0.35)
    fig.tight_layout()
    fig.savefig(out_path, dpi=160)
    plt.close(fig)
    return True


def _theta_diff_table(df: pd.DataFrame, scale: float = 1.0) -> pd.DataFrame:
    """
    Build per-theta paired awake-deep S summary.
    """
    rows = []
    for theta, subdf in df.groupby('theta'):
        sub_mean = (
            subdf.groupby(['subject', 'session'], as_index=False)['S']
            .mean()
        )
        piv = sub_mean.pivot(index='subject', columns='session', values='S')
        if not {'awake', 'deep'}.issubset(set(piv.columns)):
            continue
        paired = piv[['awake', 'deep']].dropna()
        if paired.empty:
            continue
        a = paired['awake'].to_numpy(dtype=float) * scale
        d = paired['deep'].to_numpy(dtype=float) * scale
        diff = a - d
        n = int(np.isfinite(diff).sum())
        sem = float(np.nanstd(diff, ddof=1) / np.sqrt(n)) if n > 1 else np.nan
        rows.append(
            {
                'theta': float(theta),
                'mean_diff': float(np.nanmean(diff)),
                'sem_diff': sem,
                'n': n,
            }
        )
    return pd.DataFrame(rows)


def _save_threshold_difference_plot(
    df: pd.DataFrame,
    out_path: str,
    scale: float = 1.0,
) -> float:
    """
    Save θ-scan plot and return θ* with largest |awake-deep mean difference|.
    """
    ttab = _theta_diff_table(df, scale=scale)
    if ttab.empty:
        return np.nan

    ttab = ttab.sort_values('theta')
    thetas = ttab['theta'].to_numpy(dtype=float)
    means = ttab['mean_diff'].to_numpy(dtype=float)
    sems = ttab['sem_diff'].to_numpy(dtype=float)
    i_star = int(np.nanargmax(np.abs(means)))
    theta_star = float(thetas[i_star])

    fig, ax = plt.subplots(figsize=(6.0, 4.0))
    ax.errorbar(thetas, means, yerr=sems, marker='o')
    ax.axvline(theta_star, linestyle='--')
    ax.set(xlabel='θ', ylabel='Mean S_awake–S_deep')
    fig.tight_layout()
    fig.savefig(out_path, dpi=160)
    plt.close(fig)
    return theta_star


# ---------- main API ----------

def create_doc(
    path: str,
    df: pd.DataFrame,
    df_stats_by_theta: pd.DataFrame,
    motion: dict,
    atlas_results: dict,
    mc: dict,
    theta_results: Optional[dict] = None,
    fig_dir: Optional[str] = None,
    use_sem: bool = True,
    S_SCALE: float = 1e3,
    RAM_SCALE: float = 1e3,
    CI_SCALE: float = 1.0,      # <-- NEW: scale applied to CI for display
    LABEL_S: Optional[str] = None, # if None, auto from S_SCALE
    LABEL_RAM: Optional[str] = None,
    LABEL_CI: Optional[str] = None,
    repl: Optional[Union[dict, pd.DataFrame]] = None,
):
    """
    Build the Word report directly from the long 'df' produced by compute_synergy_ci
    and the other pipeline outputs. No external df_stats bundle required.

    Parameters
    ----------
    df : DataFrame with columns:
        ['subject','session','theta','S','CI','RAM','PDI','NAS','IIM','SRPI']
        (RAM/PDI/NAS/IIM/SRPI repeat across theta → we will average per subject×session)
    df_stats_by_theta : DataFrame indexed by theta with columns ['t_S','p_S','d_S']
    motion : dict returned by motion_covariate_analysis(...)
    atlas_results : dict from atlas_check(...)
    mc : dict from compare_models(...)
    theta_results : dict(theta -> {'auc_S','p_S','auc_CI','p_CI'}), optional
    fig_dir : folder with figures like 'supp_theta_curve.png' (and optionally 'ram_curve.png')
    """

  # ---------- derive per-subject/session means ----------
    agg_map = dict(S=('S', 'mean'))
    for metric in ('CI', 'RAM', 'PDI', 'NAS', 'IIM', 'SRPI', 'IIM_raw', 'IIM_raw_scaled'):
        if metric in df.columns:
            agg_map[metric] = (metric, 'mean')
    agg = (df.groupby(['subject', 'session']).agg(**agg_map).reset_index())

    def metric_block(metric, scale=1.0):
        if metric not in agg.columns:
            return {
                'n': 0,
                'mean_a': np.nan, 'sd_a': np.nan, 'sem_a': np.nan,
                'mean_d': np.nan, 'sd_d': np.nan, 'sem_d': np.nan,
                't': np.nan, 'p': np.nan, 'd': np.nan, 'df': 0
            }
        a = agg.loc[agg.session=='awake', metric] * scale
        d = agg.loc[agg.session=='deep',  metric] * scale
        n_a, mean_a, sd_a, sem_a = _group_descriptives(a)
        n_d, mean_d, sd_d, sem_d = _group_descriptives(d)
        piv = agg.pivot(index='subject', columns='session', values=metric)
        if {'awake', 'deep'}.issubset(set(piv.columns)):
            t, p, d_eff, dfree = _paired_stats(piv['awake'], piv['deep'])
        else:
            t, p, d_eff, dfree = (np.nan, np.nan, np.nan, 0)
        return {
            'n': min(n_a, n_d),
            'mean_a': mean_a, 'sd_a': sd_a, 'sem_a': sem_a,
            'mean_d': mean_d, 'sd_d': sd_d, 'sem_d': sem_d,
            't': t, 'p': p, 'd': d_eff, 'df': dfree
        }

    S_stats   = metric_block('S',   S_SCALE)
    CI_stats  = metric_block('CI',  CI_SCALE) if 'CI' in agg.columns else None
    RAM_stats = metric_block('RAM', RAM_SCALE) if 'RAM' in agg.columns else None

    err_label = "SEM" if use_sem else "SD"
    pick_err  = (lambda stats, which: stats['sem_'+which] if use_sem else stats['sd_'+which])

    # labels (auto if not passed)
    def _label_from_scale(name, scale):
        if scale == 1.0:
            return name
        # build like "name×10³" if scale is a pure power of 10, else "name×<scale>"
        try:
            exp = int(round(np.log10(scale)))
            if np.isclose(scale, 10**exp):
                return f"{name}×10^{exp}".replace('^', '⁰¹²³⁴⁵⁶⁷⁸⁹'[exp] if 0 <= exp <= 9 else f"^{exp}")
        except Exception:
            pass
        return f"{name}×{scale:g}"

    label_S   = LABEL_S   or _label_from_scale("S",   S_SCALE)
    label_RAM = LABEL_RAM or _label_from_scale("RAM", RAM_SCALE)
    label_CI  = LABEL_CI  or _label_from_scale("CI",  CI_SCALE)

    theta_star = _best_theta_from_stats(df_stats_by_theta)
    if not np.isfinite(theta_star):
        theta_tbl = _theta_diff_table(df, scale=1.0)
        if not theta_tbl.empty:
            arr_m = theta_tbl["mean_diff"].to_numpy(dtype=float)
            if np.isfinite(arr_m).any():
                theta_arr = theta_tbl["theta"].to_numpy(dtype=float)
                theta_star = float(theta_arr[int(np.nanargmax(np.abs(arr_m)))])

    # Generate per-metric session plots for the report when a figure directory
    # is provided. These are data-driven and use subject-level paired values.
    if fig_dir:
        os.makedirs(fig_dir, exist_ok=True)
        theta_fig_path = os.path.join(fig_dir, 'supp_theta_curve.png')
        theta_star_from_fig = _save_threshold_difference_plot(df, theta_fig_path, scale=1.0)
        if np.isfinite(theta_star_from_fig):
            theta_star = float(theta_star_from_fig)

        agg_s = agg
        if 'theta' in df.columns and np.isfinite(theta_star):
            theta_vals = np.sort(df['theta'].dropna().astype(float).unique())
            if theta_vals.size:
                theta_use = float(theta_vals[np.argmin(np.abs(theta_vals - theta_star))])
                sub_s = df[np.isclose(df['theta'].astype(float), theta_use)]
                if not sub_s.empty:
                    agg_s = (
                        sub_s.groupby(['subject', 'session'], as_index=False)
                        .agg(S=('S', 'mean'))
                    )

        metric_plot_specs = []
        if 'S' in agg.columns:
            metric_plot_specs.append(("S", "s_curve.png", "Synergy (S)", label_S, S_SCALE))
        if 'CI' in agg.columns:
            metric_plot_specs.append(("CI", "ci_curve.png", "Consciousness Index (CI)", label_CI, CI_SCALE))
        if 'RAM' in agg.columns:
            metric_plot_specs.append(("RAM", "ram_curve.png", "Responsiveness-Adaptation Metric (RAM)", label_RAM, RAM_SCALE))
        if 'PDI' in agg.columns:
            metric_plot_specs.append(("PDI", "pdi_curve.png", "Phenomenal Differentiation Index (PDI)", "PDI", 1.0))
        if 'NAS' in agg.columns:
            metric_plot_specs.append(("NAS", "nas_curve.png", "Network Activation Synchrony (NAS)", "NAS", 1.0))
        if 'SRPI' in agg.columns:
            metric_plot_specs.append(("SRPI", "srpi_curve.png", "Self-Referential Processing Index (SRPI)", "SRPI", 1.0))
        # Use raw signed IIM (native units) for plots to preserve directionality/magnitude.
        if 'IIM_raw' in agg.columns:
            metric_plot_specs.append(
                ("IIM_raw", "iim_curve.png", "Integrated Information Metric (IIM raw signed)", "IIM raw (signed, unitless)", 1.0)
            )
        elif 'IIM_raw_scaled' in agg.columns:
            # Backward-compatible fallback for legacy tables.
            metric_plot_specs.append(
                ("IIM_raw_scaled", "iim_curve.png", "Integrated Information Metric (IIM raw signed)", "IIM raw (signed, unitless)", 1.0)
            )
        else:
            metric_plot_specs.append(
                ("IIM", "iim_curve.png", "Integrated Information Metric (IIM canonical)", "IIM canonical (0-1)", 1.0)
            )

        for metric, filename, title, ylabel, scale in metric_plot_specs:
            agg_src = agg_s if metric == "S" else agg
            _save_paired_metric_plot(
                agg=agg_src,
                metric=metric,
                out_path=os.path.join(fig_dir, filename),
                title=title,
                y_label=ylabel,
                scale=scale,
                sessions=('awake', 'deep'),
            )

    # % difference for S (scaled)
    pct_all = safe_pct(S_stats['mean_a'] - S_stats['mean_d'], S_stats['mean_d'])

    row_star = _safe_row_by_theta(df_stats_by_theta, theta_star) if np.isfinite(theta_star) else None

    def means_err_for_theta(theta, scale=S_SCALE):
        sub = df[np.isclose(df['theta'], theta)]
        theta_used = theta
        if sub.empty and 'theta' in df.columns:
            theta_vals = np.sort(df['theta'].dropna().astype(float).unique())
            if theta_vals.size:
                theta_used = float(theta_vals[np.argmin(np.abs(theta_vals - theta))])
                sub = df[np.isclose(df['theta'], theta_used)]
        a = sub.loc[sub.session=='awake', 'S'] * scale
        d = sub.loc[sub.session=='deep',  'S'] * scale
        _, ma, sda, sema = _group_descriptives(a)
        _, md, sdd, semd = _group_descriptives(d)
        ea = sema if use_sem else sda
        ed = semd if use_sem else sdd
        return theta_used, ma, ea, md, ed

    if np.isfinite(theta_star):
        theta_used, m_star_a, e_star_a, m_star_d, e_star_d = means_err_for_theta(float(theta_star))
        pct_star = safe_pct(m_star_a - m_star_d, m_star_d)
    else:
        theta_used = np.nan
        m_star_a = e_star_a = m_star_d = e_star_d = np.nan
        pct_star = "na"

    # ---------- build document ----------
    doc = Document()
    doc.add_heading("Empirical Validation of Synergy & MPC Metrics in fMRI", level=1)

    # 1) Main effects: Synergy
    theta_detail = (
        f"\nThreshold-specific analysis at θ*={fmt(theta_used,decimals=2)} "
        f"(largest |awake−deep| mean S difference): "
        f"awake {fmt(m_star_a)}±{fmt(e_star_a)} ({err_label}) vs "
        f"deep {fmt(m_star_d)}±{fmt(e_star_d)} ({err_label}); "
        f"t={fmt(getattr(row_star,'t_S',np.nan))}, "
        f"p={fmt(getattr(row_star,'p_S',np.nan),decimals=4)}, "
        f"d={fmt(getattr(row_star,'d_S',np.nan))}; Δ={pct_star}."
        if np.isfinite(theta_used)
        else "\nThreshold-specific analysis: θ* could not be identified from available data."
    )

    doc.add_paragraph(
        f"In n={S_stats['n']} subjects, average {label_S} differed by {pct_all} between awake "
        f"({fmt(S_stats['mean_a'])} ± {fmt(pick_err(S_stats,'a'))} [{err_label}]) and deep "
        f"({fmt(S_stats['mean_d'])} ± {fmt(pick_err(S_stats,'d'))} [{err_label}]); "
        f"t({S_stats['df']})={fmt(S_stats['t'])}, p={fmt(S_stats['p'], decimals=4)}, d={fmt(S_stats['d'])}."
        + theta_detail
    )

    # 2) CI and RAM (if present in this run)
    if CI_stats is not None:
        doc.add_paragraph(
            f"{label_CI}: awake {fmt(CI_stats['mean_a'])} ± {fmt(pick_err(CI_stats,'a'))} ({err_label}) vs "
            f"deep {fmt(CI_stats['mean_d'])} ± {fmt(pick_err(CI_stats,'d'))} ({err_label}); "
            f"t({CI_stats['df']})={fmt(CI_stats['t'])}, p={fmt(CI_stats['p'],decimals=4)}, d={fmt(CI_stats['d'])}."
        )
    if RAM_stats is not None:
        doc.add_paragraph(
            f"{label_RAM}: awake {fmt(RAM_stats['mean_a'])} ± {fmt(pick_err(RAM_stats,'a'))} ({err_label}) vs "
            f"deep {fmt(RAM_stats['mean_d'])} ± {fmt(pick_err(RAM_stats,'d'))} ({err_label}); "
            f"t({RAM_stats['df']})={fmt(RAM_stats['t'])}, p={fmt(RAM_stats['p'],decimals=4)}, d={fmt(RAM_stats['d'])}."
        )

    # 3) MPC components snapshot (subset-aware)
    metric_labels = {
        'PDI': 'PDI',
        'NAS': 'NAS',
        'IIM': 'IIM canonical (0-1)',
        'SRPI': 'SRPI',
    }
    metric_order = [m for m in ('PDI', 'NAS', 'IIM', 'SRPI') if m in agg.columns]
    for metric in metric_order:
        label = metric_labels.get(metric, metric)
        st = metric_block(metric, scale=1.0)
        doc.add_paragraph(
            f"{label}: awake {fmt(st['mean_a'],decimals=4)} ± {fmt(pick_err(st,'a'),decimals=4)} ({err_label}) vs "
            f"deep {fmt(st['mean_d'],decimals=4)} ± {fmt(pick_err(st,'d'),decimals=4)} ({err_label}); "
            f"t({st['df']})={fmt(st['t'])}, p={fmt(st['p'],decimals=4)}, d={fmt(st['d'])}."
        )
    if 'IIM_raw' in agg.columns:
        st_raw = metric_block('IIM_raw', scale=1.0)
        doc.add_paragraph(
            f"IIM raw (signed, unitless): awake {fmt(st_raw['mean_a'],decimals=4)} ± "
            f"{fmt(pick_err(st_raw,'a'),decimals=4)} ({err_label}) vs "
            f"deep {fmt(st_raw['mean_d'],decimals=4)} ± {fmt(pick_err(st_raw,'d'),decimals=4)} ({err_label}); "
            f"t({st_raw['df']})={fmt(st_raw['t'])}, p={fmt(st_raw['p'],decimals=4)}, d={fmt(st_raw['d'])}."
        )
    elif 'IIM_raw_scaled' in agg.columns:
        st_raw = metric_block('IIM_raw_scaled', scale=1.0)
        doc.add_paragraph(
            f"IIM raw (signed, unitless): awake {fmt(st_raw['mean_a'],decimals=4)} ± "
            f"{fmt(pick_err(st_raw,'a'),decimals=4)} ({err_label}) vs "
            f"deep {fmt(st_raw['mean_d'],decimals=4)} ± {fmt(pick_err(st_raw,'d'),decimals=4)} ({err_label}); "
            f"t({st_raw['df']})={fmt(st_raw['t'])}, p={fmt(st_raw['p'],decimals=4)}, d={fmt(st_raw['d'])}."
        )

    # 4) Motion
    if isinstance(motion, dict):
        if 'skipped' in motion:
            doc.add_paragraph(f"Motion covariate analysis: skipped ({motion['skipped']})")
        else:
            coefs = ", ".join(fmt(c) for c in motion.get('coef_awake', []))
            pvals = ", ".join(fmt(p, decimals=4) for p in motion.get('p_awake', []))
            doc.add_paragraph(f"Motion covariates (awake): coefs=[{coefs}]; p-values=[{pvals}].")

    # 5) Atlas robustness
    if isinstance(atlas_results, dict):
        for name, r in atlas_results.items():
            if isinstance(r, dict) and ("roughness" in r):
                rough = r.get("roughness") or {}
                if isinstance(rough, dict) and rough:
                    rough_items = ", ".join(
                        f"{k}={fmt(v,decimals=3)}" for k, v in rough.items()
                    )
                    doc.add_paragraph(f"Atlas {name}: S roughness {rough_items}.")
                else:
                    doc.add_paragraph(f"Atlas {name}: S roughness not available.")

                metrics = r.get("metrics") if isinstance(r.get("metrics"), dict) else {}
                metric_order = ("S", "PDI", "NAS", "IIM", "IIM_raw", "CI", "RAM", "SRPI")
                for m in metric_order:
                    st = metrics.get(m)
                    if not isinstance(st, dict):
                        continue
                    n = int(st.get("n") or 0)
                    if n <= 0:
                        continue
                    doc.add_paragraph(
                        f"Atlas {name} {m}: n={n}; "
                        f"mean_awake={fmt(st.get('mean_a', np.nan),decimals=4)}, "
                        f"mean_deep={fmt(st.get('mean_b', np.nan),decimals=4)}, "
                        f"Δ(awake−deep)={fmt(st.get('delta_a_minus_b', np.nan),decimals=4)}, "
                        f"t={fmt(st.get('t', np.nan))}, "
                        f"p={fmt(st.get('p', np.nan),decimals=4)}, "
                        f"d={fmt(st.get('d_paired', np.nan))}."
                    )
                notes = r.get("notes")
                if isinstance(notes, list) and notes:
                    doc.add_paragraph(f"Atlas {name} notes: {' '.join(str(x) for x in notes)}")
            else:
                try:
                    doc.add_paragraph(f"Atlas {name}: metric roughness awake={fmt(r['awake'],decimals=3)}, deep={fmt(r['deep'],decimals=3)}.")
                except Exception:
                    doc.add_paragraph(f"Atlas {name}: {r}")

    # 6) Synergy by θ
    if df_stats_by_theta is not None and not df_stats_by_theta.empty:
        doc.add_heading("Synergy by θ", level=2)
        for theta, row in df_stats_by_theta.sort_index().iterrows():
            mean_diff_str = ""
            if "mean_diff_S" in df_stats_by_theta.columns:
                mean_diff_str = f", mean_diff_S={fmt(getattr(row, 'mean_diff_S', np.nan))}"
            doc.add_paragraph(
                f"θ={fmt(theta,decimals=2)}: t_S={fmt(row.t_S)}, p_S={fmt(row.p_S,decimals=4)}, "
                f"d_S={fmt(row.d_S)}{mean_diff_str}"
            )

    # 7) Permutation-Test AUC by θ
    if isinstance(theta_results, dict) and len(theta_results):
        doc.add_heading("Permutation-Test AUC by θ", level=2)
        for theta in sorted(theta_results):
            res = theta_results[theta]
            doc.add_paragraph(
                f"θ={fmt(theta,decimals=2)}: AUC_S={fmt(res.get('auc_S',np.nan),decimals=3)}, "
                f"p_S={fmt(res.get('p_S',np.nan),decimals=4)}; "
                f"AUC_CI={fmt(res.get('auc_CI',np.nan),decimals=3)}, "
                f"p_CI={fmt(res.get('p_CI',np.nan),decimals=4)}"
            )

    # 8) Model comparisons
    if isinstance(mc, dict):
        for m, res in mc.items():
            try:
                doc.add_paragraph(f"Model {m}: ΔAUC={fmt(res['delta_auc'],decimals=3)}, p={fmt(res['p_val'],decimals=3)}.")
            except Exception:
                doc.add_paragraph(f"Model {m}: {res}")
                
    # 8b) Independent replication
    if repl is not None:
        doc.add_heading("Independent Replication", level=2)

        def _write_rep_row(metric_name, stats_tuple):
            t, p, d_eff, dfree, mean_a, mean_d, err_a, err_d, n = stats_tuple
            doc.add_paragraph(
                f"{metric_name}: n={n}; awake {mean_a:.4f}±{err_a:.4f} ({err_label}) vs "
                f"deep {mean_d:.4f}±{err_d:.4f} ({err_label}); "
                f"t({dfree})={t:.2f}, p={p:.4f}, d={d_eff:.2f}."
            )

        if isinstance(repl, dict):
            # Try a compact dict schema; fallbacks are safe-printed if keys differ
            dataset = repl.get('dataset', 'independent dataset')
            doc.add_paragraph(f"Dataset: {dataset}")
            if {'delta_S','ci','cohend'}.issubset(repl.keys()):
                lo, hi = repl['ci'] if isinstance(repl['ci'], (list, tuple)) and len(repl['ci'])==2 else (np.nan, np.nan)
                doc.add_paragraph(
                    f"ΔS={repl['delta_S']:.3f} (95% CI [{lo:.3f}–{hi:.3f}]), "
                    f"Cohen’s d={repl['cohend']:.2f}."
                )
            else:
                doc.add_paragraph(f"(Replication summary keys not recognized — got: {list(repl.keys())})")

        elif isinstance(repl, pd.DataFrame):
            # Compute paired stats on replication as we did for the main set
            rep = (repl.groupby(['subject','session'])
                        .agg(S=('S','mean'), CI=('CI','mean'))
                        .reset_index())

            def _paired_block(metric):
                piv = rep.pivot(index='subject', columns='session', values=metric)
                t, p, d_eff, dfree = _paired_stats(piv['awake'], piv['deep'])
                a = rep.loc[rep.session=='awake', metric]
                d = rep.loc[rep.session=='deep',  metric]
                n_a, mean_a, sd_a, sem_a = _group_descriptives(a)
                n_d, mean_d, sd_d, sem_d = _group_descriptives(d)
                n = min(n_a, n_d)
                err_a = sem_a if use_sem else sd_a
                err_d = sem_d if use_sem else sd_d
                return t, p, d_eff, dfree, mean_a, mean_d, err_a, err_d, n

            _write_rep_row('Synergy (S)', _paired_block('S'))
            _write_rep_row('Consciousness Index (CI)', _paired_block('CI'))

        else:
            doc.add_paragraph(f"(Unsupported replication object type: {type(repl)})")


    # 9) Figures
    if fig_dir:
        theta_fig = os.path.join(fig_dir, 'supp_theta_curve.png')
        if os.path.exists(theta_fig):
            doc.add_heading("Supplementary: Synergy difference across θ", level=2)
            doc.add_picture(theta_fig, width=Inches(5.0))
        else:
            doc.add_paragraph(f"(Could not find θ-curve at {theta_fig})")

        figure_specs = []
        if 'S' in agg.columns:
            figure_specs.append(("s_curve.png", "Supplementary: Synergy (S)"))
        if 'CI' in agg.columns:
            figure_specs.append(("ci_curve.png", "Supplementary: Consciousness Index (CI)"))
        if 'RAM' in agg.columns:
            figure_specs.append(("ram_curve.png", "Supplementary: Responsiveness-Adaptation Metric (RAM)"))
        if 'PDI' in agg.columns:
            figure_specs.append(("pdi_curve.png", "Supplementary: Phenomenal Differentiation Index (PDI)"))
        if 'NAS' in agg.columns:
            figure_specs.append(("nas_curve.png", "Supplementary: Network Activation Synchrony (NAS)"))
        if ('IIM_raw_scaled' in agg.columns) or ('IIM' in agg.columns):
            figure_specs.append(("iim_curve.png", "Supplementary: Integrated Information Metric (IIM raw signed)"))
        if 'SRPI' in agg.columns:
            figure_specs.append(("srpi_curve.png", "Supplementary: Self-Referential Processing Index (SRPI)"))
        for filename, heading in figure_specs:
            fig_path = os.path.join(fig_dir, filename)
            if os.path.exists(fig_path):
                doc.add_heading(heading, level=2)
                doc.add_picture(fig_path, width=Inches(5.0))

    doc.save(path)
